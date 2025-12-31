"""
PDF Toolkit Tool

This tool provides comprehensive PDF manipulation capabilities.
Features:
- PDF merging and splitting
- PDF compression
- Add watermarks to PDFs
- Convert PDF to images
- Extract text from PDF
- Batch processing capabilities
"""

import os
import tempfile
from io import BytesIO
from flask import Blueprint, render_template, request, jsonify, send_file, current_app
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader, PdfWriter, PdfMerger
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.colors import grey
import fitz  # PyMuPDF for PDF to image conversion
from PIL import Image
import zipfile
from datetime import datetime

pdf_toolkit_bp = Blueprint('pdf_toolkit', __name__,
                         template_folder='templates',
                         url_prefix='/pdf-toolkit')

def allowed_file(filename):
    """Check if the file has an allowed extension."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() == 'pdf'

def compress_pdf(input_path, output_path, quality=0.8):
    """Compress a PDF by rasterizing pages to images and re-encoding with lower JPEG quality.

    Note: This approach rasterizes PDF pages which may increase file size for text-heavy PDFs
    but effectively reduces size for image-rich PDFs. `quality` expected 1-100 (higher = better quality).
    """
    try:
        q = int(float(quality))
        if q <= 0:
            q = 75
        if q > 100:
            q = 100

        # Rasterize pages to images then recompress to JPEG and reassemble
        with tempfile.TemporaryDirectory() as tmpdir:
            img_dir = os.path.join(tmpdir, 'images')
            os.makedirs(img_dir, exist_ok=True)

            images = pdf_to_images(input_path, img_dir, dpi=150, format='PNG')
            if not images:
                return False

            jpg_paths = []
            for img_path in images:
                try:
                    im = Image.open(img_path).convert('RGB')
                    jpg_path = os.path.splitext(img_path)[0] + '.jpg'
                    im.save(jpg_path, 'JPEG', quality=q, optimize=True)
                    jpg_paths.append(jpg_path)
                except Exception:
                    # fallback: keep original PNG if conversion fails
                    jpg_paths.append(img_path)

            # Assemble into PDF
            return images_to_pdf(jpg_paths, output_path)
    except Exception as e:
        print(f"Compression error: {e}")
        return False

def merge_pdfs(pdf_paths, output_path):
    """Merge multiple PDF files into one."""
    try:
        merger = PdfMerger()

        for pdf_path in pdf_paths:
            merger.append(pdf_path)

        merger.write(output_path)
        merger.close()

        return True
    except Exception as e:
        print(f"Merge error: {e}")
        return False

def split_pdf(input_path, output_dir, split_type='page', pages_per_file=1, custom_ranges=None):
    """Split a PDF file into multiple files.

    Supported split_type values:
    - 'page' : each page -> one file
    - 'pages_per_pdf' or 'range' : split into chunks of `pages_per_file`
    - 'even_odd' : two files, odd pages and even pages
    - 'halve' : split into two roughly equal halves
    - 'custom' : use `custom_ranges` string like '1-3,5,7-9' to create files
    """
    try:
        reader = PdfReader(input_path)
        total_pages = len(reader.pages)
        split_files = []

        def _write_writer_if_has_pages(w, name):
            if len(w.pages) > 0:
                output_file = os.path.join(output_dir, name)
                with open(output_file, 'wb') as f:
                    w.write(f)
                split_files.append(output_file)

        if split_type == 'page':
            for i, page in enumerate(reader.pages):
                writer = PdfWriter()
                writer.add_page(page)
                _write_writer_if_has_pages(writer, f'page_{i+1}.pdf')

        elif split_type in ('pages_per_pdf', 'range'):
            # Split into chunks of pages_per_file
            if pages_per_file <= 0:
                pages_per_file = 1

            for i in range(0, total_pages, pages_per_file):
                writer = PdfWriter()
                end_page = min(i + pages_per_file, total_pages)
                for j in range(i, end_page):
                    writer.add_page(reader.pages[j])
                _write_writer_if_has_pages(writer, f'pages_{i+1}-{end_page}.pdf')

        elif split_type == 'even_odd':
            odd_writer = PdfWriter()
            even_writer = PdfWriter()
            for idx, page in enumerate(reader.pages):
                # human page numbers: 1 -> odd, 2 -> even
                if (idx % 2) == 0:
                    odd_writer.add_page(page)
                else:
                    even_writer.add_page(page)

            _write_writer_if_has_pages(odd_writer, 'pages_odd.pdf')
            _write_writer_if_has_pages(even_writer, 'pages_even.pdf')

        elif split_type == 'halve':
            # Split into two halves
            mid = (total_pages + 1) // 2
            first_writer = PdfWriter()
            second_writer = PdfWriter()

            for j in range(0, mid):
                first_writer.add_page(reader.pages[j])
            for j in range(mid, total_pages):
                second_writer.add_page(reader.pages[j])

            _write_writer_if_has_pages(first_writer, f'pages_1-{mid}.pdf')
            _write_writer_if_has_pages(second_writer, f'pages_{mid+1}-{total_pages}.pdf')

        elif split_type == 'custom':
            # custom_ranges expected like '1-3,5,7-9'
            ranges_str = (custom_ranges or '').strip()
            if not ranges_str:
                return []

            parts = [p.strip() for p in ranges_str.split(',') if p.strip()]
            for part in parts:
                if '-' in part:
                    start_s, end_s = part.split('-', 1)
                    start = max(1, int(start_s.strip()))
                    end = min(total_pages, int(end_s.strip()))
                else:
                    start = end = int(part)

                # convert to 0-based indexes
                start_idx = max(0, start - 1)
                end_idx = min(total_pages - 1, end - 1)
                if start_idx > end_idx:
                    continue

                writer = PdfWriter()
                for j in range(start_idx, end_idx + 1):
                    writer.add_page(reader.pages[j])

                _write_writer_if_has_pages(writer, f'pages_{start}-{end}.pdf')

        else:
            # Unknown split type
            return []

        return split_files
    except Exception as e:
        print(f"Split error: {e}")
        return []

def add_watermark(input_path, output_path, watermark_text, opacity=0.3):
    """Add a text watermark to a PDF file."""
    try:
        reader = PdfReader(input_path)
        writer = PdfWriter()

        for page in reader.pages:
            # Create watermark PDF
            watermark_packet = BytesIO()
            can = canvas.Canvas(watermark_packet, pagesize=letter)
            can.setFont("Helvetica", 40)
            can.setFillColorRGB(0.5, 0.5, 0.5, alpha=opacity)
            can.saveState()
            can.translate(300, 300)
            can.rotate(45)
            can.drawString(0, 0, watermark_text)
            can.restoreState()
            can.save()

            watermark_packet.seek(0)
            watermark = PdfReader(watermark_packet)
            watermark_page = watermark.pages[0]

            # Merge watermark with page
            page.merge_page(watermark_page)
            writer.add_page(page)

        with open(output_path, 'wb') as f:
            writer.write(f)

        return True
    except Exception as e:
        print(f"Watermark error: {e}")
        return False

def pdf_to_images(input_path, output_dir, dpi=150, format='PNG'):
    """Convert PDF pages to images."""
    try:
        doc = fitz.open(input_path)
        image_files = []

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            pix = page.get_pixmap(matrix=fitz.Matrix(dpi/72, dpi/72))

            output_file = os.path.join(output_dir, f'page_{page_num+1}.{format.lower()}')
            pix.save(output_file)
            image_files.append(output_file)

        doc.close()
        return image_files
    except Exception as e:
        print(f"PDF to images error: {e}")
        return []

# Text extraction removed — PDF Toolkit no longer exposes text extraction via the web UI.

def rotate_pdf(input_path, output_path, angle=90):
    """Rotate all pages in a PDF."""
    try:
        reader = PdfReader(input_path)
        writer = PdfWriter()

        # Normalize angle to 0, 90, 180, 270
        angle = angle % 360

        for page in reader.pages:
            if angle != 0:
                page.rotate(angle)
            writer.add_page(page)

        with open(output_path, 'wb') as f:
            writer.write(f)

        return True
    except Exception as e:
        print(f"Rotation error: {e}")
        return False

def remove_pages_from_pdf(input_path, output_path, pages_to_remove):
    """Remove specific pages from a PDF."""
    try:
        reader = PdfReader(input_path)
        writer = PdfWriter()

        # Parse pages_to_remove (e.g., "1,3,5" or "1-5")
        pages_set = set()
        
        for part in pages_to_remove.split(','):
            part = part.strip()
            if '-' in part:
                start, end = part.split('-')
                pages_set.update(range(int(start.strip())-1, int(end.strip())))
            else:
                pages_set.add(int(part)-1)

        for idx, page in enumerate(reader.pages):
            if idx not in pages_set:
                writer.add_page(page)

        with open(output_path, 'wb') as f:
            writer.write(f)

        return True
    except Exception as e:
        print(f"Remove pages error: {e}")
        return False

def images_to_pdf(image_paths, output_path):
    """Convert multiple images to a single PDF."""
    try:
        from PIL import Image
        
        images = []
        for img_path in sorted(image_paths):
            img = Image.open(img_path)
            # Convert RGBA to RGB if necessary
            if img.mode in ('RGBA', 'LA', 'P'):
                background = Image.new('RGB', img.size, (255, 255, 255))
                background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                img = background
            images.append(img.convert('RGB'))

        if images:
            images[0].save(output_path, 'PDF', save_all=True, append_images=images[1:])
            return True
        
        return False
    except Exception as e:
        print(f"Images to PDF error: {e}")
        return False

def image_to_single_pdf(image_path, output_path):
    """Convert a single image to a single-page PDF."""
    try:
        img = Image.open(image_path)
        if img.mode in ('RGBA', 'LA', 'P'):
            background = Image.new('RGB', img.size, (255, 255, 255))
            background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
            img = background
        img.convert('RGB').save(output_path, 'PDF')
        return True
    except Exception as e:
        print(f"Image to single PDF error: {e}")
        return False

def pdf_to_word(input_path, output_path):
    """Convert PDF to Word document."""
    try:
        # Try using pdfplumber for better text extraction
        try:
            import pdfplumber
            use_pdfplumber = True
        except ImportError:
            use_pdfplumber = False

        try:
            from docx import Document
        except ImportError:
            print("python-docx not installed")
            return False

        doc = Document()
        
        if use_pdfplumber:
            with pdfplumber.open(input_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text from page
                    text = page.extract_text()
                    if text:
                        doc.add_paragraph(text)
                    
                    # Add page break except for last page
                    if page_num < len(pdf.pages):
                        doc.add_page_break()
        else:
            # Fallback to PyPDF2 if pdfplumber is not available
            reader = PdfReader(input_path)
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text:
                    doc.add_paragraph(text)
                
                # Add page break except for last page
                if page_num < len(reader.pages):
                    doc.add_page_break()

        doc.save(output_path)
        return True
    except Exception as e:
        print(f"PDF to Word error: {e}")
        return False

@pdf_toolkit_bp.route('/')
def index():
    """Main page for the PDF Toolkit."""
    return render_template('tools/pdf_toolkit.html')

@pdf_toolkit_bp.route('/process', methods=['POST'])
def process_pdfs():
    """Handle PDF processing operations."""
    try:
        if 'pdfs' not in request.files:
            return jsonify({'error': 'No files uploaded'}), 400

        files = request.files.getlist('pdfs')
        operation = request.form.get('operation')

        if not files or files[0].filename == '':
            return jsonify({'error': 'No files selected'}), 400

        # Create temporary directory for processing
        with tempfile.TemporaryDirectory() as temp_dir:
            # Allowed image extensions
            image_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'}

            # For images_to_pdf, accept only images.
            # For merge, accept PDFs and images (images will be converted to single-page PDFs).
            if operation == 'images_to_pdf':
                uploaded_files = [file for file in files if file and os.path.splitext(file.filename)[1].lower() in image_extensions]
                if not uploaded_files:
                    return jsonify({'error': 'No valid image files found'}), 400
            elif operation == 'merge':
                allowed_exts = set(image_extensions)
                allowed_exts.add('.pdf')
                uploaded_files = [file for file in files if file and os.path.splitext(file.filename)[1].lower() in allowed_exts]
                if not uploaded_files:
                    return jsonify({'error': 'No valid PDF or image files found for merging'}), 400
            else:
                uploaded_files = [file for file in files if file and allowed_file(file.filename)]
                if not uploaded_files:
                    return jsonify({'error': 'No valid PDF files found'}), 400

            # Save uploaded files
            uploaded_paths = []
            for file in uploaded_files:
                filename = secure_filename(file.filename)
                file_path = os.path.join(temp_dir, filename)
                file.save(file_path)
                uploaded_paths.append(file_path)

            # Process based on operation
            result_file = None
            result_files = []
            extracted_text = None

            try:
                if operation == 'merge':
                    if len(uploaded_paths) < 1:
                        return jsonify({'error': 'At least 1 file required for merging'}), 400

                    # Prepare list of PDF files to merge. Convert images to single-page PDFs.
                    to_merge = []
                    for p in uploaded_paths:
                        ext = os.path.splitext(p)[1].lower()
                        if ext in image_extensions:
                            tmp_pdf = os.path.join(temp_dir, os.path.splitext(os.path.basename(p))[0] + '_img.pdf')
                            if image_to_single_pdf(p, tmp_pdf):
                                to_merge.append(tmp_pdf)
                        elif ext == '.pdf':
                            to_merge.append(p)

                    if not to_merge:
                        return jsonify({'error': 'No convertible files found for merging'}), 400

                    output_file = os.path.join(temp_dir, 'merged.pdf')
                    if merge_pdfs(to_merge, output_file):
                        result_file = output_file

                elif operation == 'split':
                    if len(uploaded_paths) != 1:
                        return jsonify({'error': 'Exactly 1 PDF file required for splitting'}), 400

                    split_type = request.form.get('split_type', 'page')
                    pages_per_file = int(request.form.get('pages_per_file', 1))
                    custom_ranges = request.form.get('custom_ranges', '')

                    output_dir = os.path.join(temp_dir, 'split')
                    os.makedirs(output_dir, exist_ok=True)

                    result_files = split_pdf(uploaded_paths[0], output_dir, split_type, pages_per_file, custom_ranges)

                elif operation == 'compress':
                    if len(uploaded_paths) != 1:
                        return jsonify({'error': 'Exactly 1 PDF file required for compression'}), 400

                    quality = float(request.form.get('quality', 0.8))
                    output_file = os.path.join(temp_dir, 'compressed.pdf')

                    if compress_pdf(uploaded_paths[0], output_file, quality):
                        result_file = output_file

                elif operation == 'watermark':
                    if len(uploaded_paths) != 1:
                        return jsonify({'error': 'Exactly 1 PDF file required for watermarking'}), 400

                    watermark_text = request.form.get('watermark_text', 'CONFIDENTIAL')
                    opacity = float(request.form.get('opacity', 0.3))
                    output_file = os.path.join(temp_dir, 'watermarked.pdf')

                    if add_watermark(uploaded_paths[0], output_file, watermark_text, opacity):
                        result_file = output_file

                elif operation == 'convert':
                    if len(uploaded_paths) != 1:
                        return jsonify({'error': 'Exactly 1 PDF file required for conversion'}), 400

                    dpi = int(request.form.get('dpi', 150))
                    output_format = request.form.get('format', 'PNG')
                    output_dir = os.path.join(temp_dir, 'images')
                    os.makedirs(output_dir, exist_ok=True)

                    result_files = pdf_to_images(uploaded_paths[0], output_dir, dpi, output_format)

                # 'extract' operation removed

                elif operation == 'rotate':
                    if len(uploaded_paths) != 1:
                        return jsonify({'error': 'Exactly 1 PDF file required for rotation'}), 400

                    angle = int(request.form.get('angle', 90))
                    output_file = os.path.join(temp_dir, 'rotated.pdf')

                    if rotate_pdf(uploaded_paths[0], output_file, angle):
                        result_file = output_file

                elif operation == 'remove_pages':
                    if len(uploaded_paths) != 1:
                        return jsonify({'error': 'Exactly 1 PDF file required for page removal'}), 400

                    pages_to_remove = request.form.get('pages', '')
                    if not pages_to_remove:
                        return jsonify({'error': 'Please specify pages to remove'}), 400

                    output_file = os.path.join(temp_dir, 'trimmed.pdf')

                    if remove_pages_from_pdf(uploaded_paths[0], output_file, pages_to_remove):
                        result_file = output_file

                elif operation == 'images_to_pdf':
                    if len(uploaded_paths) < 1:
                        return jsonify({'error': 'At least 1 image file required'}), 400

                    output_file = os.path.join(temp_dir, 'converted.pdf')

                    if images_to_pdf(uploaded_paths, output_file):
                        result_file = output_file

                elif operation == 'pdf_to_word':
                    if len(uploaded_paths) != 1:
                        return jsonify({'error': 'Exactly 1 PDF file required for Word conversion'}), 400

                    output_file = os.path.join(temp_dir, 'converted.docx')

                    if pdf_to_word(uploaded_paths[0], output_file):
                        result_file = output_file

                else:
                    return jsonify({'error': 'Invalid operation'}), 400

                # Prepare response
                # No special JSON response types remain; all successful operations return files/zip
                if result_files:
                    # Multiple files result (split, convert)
                    zip_buffer = BytesIO()
                    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                        for file_path in result_files:
                            filename = os.path.basename(file_path)
                            zip_file.write(file_path, filename)

                    zip_buffer.seek(0)

                    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    filename = f"pdf_{operation}_{timestamp}.zip"

                    return send_file(
                        zip_buffer,
                        as_attachment=True,
                        download_name=filename,
                        mimetype='application/zip'
                    )

                elif result_file and os.path.exists(result_file):
                    # Single file result
                    with open(result_file, 'rb') as f:
                        file_data = f.read()

                    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    filename = f"pdf_{operation}_{timestamp}.pdf"

                    return send_file(
                        BytesIO(file_data),
                        as_attachment=True,
                        download_name=filename,
                        mimetype='application/pdf'
                    )
                else:
                    return jsonify({'error': 'Operation failed'}), 500

            except Exception as e:
                return jsonify({'error': f'Processing failed: {str(e)}'}), 500

    except Exception as e:
        return jsonify({'error': f'Server error: {str(e)}'}), 500

@pdf_toolkit_bp.route('/validate', methods=['POST'])
def validate_pdfs():
    """Validate uploaded PDF files."""
    try:
        if 'pdfs' not in request.files:
            return jsonify({'error': 'No files uploaded'}), 400

        files = request.files.getlist('pdfs')
        valid_files = []
        invalid_files = []

        for file in files:
            if file and allowed_file(file.filename):
                # Try to read the PDF to ensure it's valid
                try:
                    file_content = file.read()
                    if len(file_content) > 0:
                        file.seek(0)
                        valid_files.append({
                            'name': file.filename,
                            'size': len(file_content)
                        })
                    else:
                        invalid_files.append({'name': file.filename, 'error': 'Empty file'})
                except Exception as e:
                    invalid_files.append({'name': file.filename, 'error': str(e)})
            else:
                invalid_files.append({'name': file.filename, 'error': 'Invalid PDF file'})

        return jsonify({
            'success': True,
            'valid_files': valid_files,
            'invalid_files': invalid_files,
            'total_files': len(files)
        })

    except Exception as e:
        return jsonify({'error': f'Validation failed: {str(e)}'}), 500

# Error handlers
@pdf_toolkit_bp.errorhandler(413)
def too_large(e):
    return jsonify({'error': 'File too large'}), 413

@pdf_toolkit_bp.errorhandler(400)
def bad_request(e):
    return jsonify({'error': 'Bad request'}), 400